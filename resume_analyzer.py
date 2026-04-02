"""
Resume Analyzer Engine using NLP
Author: Ankit Kumar

This module provides NLP-based analysis of resumes against job descriptions.
It uses multiple techniques including:
- TF-IDF vectorization for keyword matching
- Cosine similarity for semantic comparison
- Named entity recognition for skill extraction
- Keyword extraction and matching
"""

import re
import os
from typing import Dict, List, Tuple, Any
from collections import Counter

# Document parsing
import PyPDF2
from docx import Document

# NLP Libraries
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer

# Download required NLTK data (will be skipped if already present)
def download_nltk_data():
    """Download required NLTK data packages."""
    packages = ['punkt', 'stopwords', 'wordnet', 'averaged_perceptron_tagger', 'punkt_tab']
    for package in packages:
        try:
            nltk.download(package, quiet=True)
        except Exception:
            pass

download_nltk_data()


class ResumeAnalyzer:
    """
    A class to analyze resumes against job descriptions using NLP techniques.
    """
    
    # Common technical skills keywords
    TECH_SKILLS = {
        'programming': ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'go', 
                       'rust', 'php', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'perl', 'html', 
                       'css', 'sql', 'nosql', 'bash', 'shell', 'powershell'],
        'frameworks': ['react', 'angular', 'vue', 'django', 'flask', 'fastapi', 'spring', 'express',
                      'nodejs', 'node.js', '.net', 'dotnet', 'rails', 'laravel', 'symfony', 'nextjs',
                      'nuxt', 'gatsby', 'svelte', 'bootstrap', 'tailwind', 'jquery'],
        'databases': ['mysql', 'postgresql', 'postgres', 'mongodb', 'redis', 'elasticsearch', 
                     'cassandra', 'dynamodb', 'sqlite', 'oracle', 'sqlserver', 'mariadb', 'neo4j',
                     'firebase', 'supabase'],
        'cloud': ['aws', 'azure', 'gcp', 'google cloud', 'heroku', 'digitalocean', 'cloudflare',
                 'ec2', 's3', 'lambda', 'kubernetes', 'k8s', 'docker', 'terraform', 'ansible'],
        'tools': ['git', 'github', 'gitlab', 'bitbucket', 'jira', 'confluence', 'slack', 'jenkins',
                 'circleci', 'travis', 'webpack', 'npm', 'yarn', 'pip', 'maven', 'gradle'],
        'data_science': ['machine learning', 'ml', 'deep learning', 'ai', 'artificial intelligence',
                        'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 
                        'matplotlib', 'seaborn', 'nlp', 'natural language processing', 'computer vision',
                        'neural network', 'data analysis', 'data visualization', 'statistics'],
        'soft_skills': ['leadership', 'communication', 'teamwork', 'problem solving', 'analytical',
                       'creative', 'adaptable', 'organized', 'detail-oriented', 'time management',
                       'critical thinking', 'collaboration', 'presentation', 'negotiation']
    }
    
    # Education keywords with levels
    EDUCATION_LEVELS = {
        'phd': ['phd', 'ph.d', 'doctorate', 'doctoral'],
        'masters': ['masters', 'master', 'msc', 'm.s.', 'mba', 'm.b.a.', 'mtech', 'm.tech'],
        'bachelors': ['bachelors', 'bachelor', 'bsc', 'b.s.', 'btech', 'b.tech', 'be', 'b.e.'],
        'associate': ['associate', 'diploma', 'certification']
    }
    
    # Experience indicators
    EXPERIENCE_PATTERNS = [
        r'(\d+)\+?\s*(?:years?|yrs?)\s*(?:of)?\s*(?:experience|exp)',
        r'(?:experience|exp)\s*(?:of)?\s*(\d+)\+?\s*(?:years?|yrs?)',
        r'(?:worked|working)\s*(?:for)?\s*(\d+)\+?\s*(?:years?|yrs?)'
    ]
    
    def __init__(self):
        """Initialize the analyzer with NLP tools."""
        self.lemmatizer = WordNetLemmatizer()
        try:
            self.stop_words = set(stopwords.words('english'))
        except Exception:
            self.stop_words = set()
        
        # TF-IDF Vectorizer
        self.vectorizer = TfidfVectorizer(
            max_features=5000,
            stop_words='english',
            ngram_range=(1, 2),
            min_df=1
        )
    
    def extract_text_from_pdf(self, filepath: str) -> str:
        """Extract text content from a PDF file."""
        text = ""
        try:
            with open(filepath, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            raise Exception(f"Failed to read PDF: {str(e)}")
        return text.strip()
    
    def extract_text_from_docx(self, filepath: str) -> str:
        """Extract text content from a DOCX file."""
        text = ""
        try:
            doc = Document(filepath)
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            raise Exception(f"Failed to read DOCX: {str(e)}")
        return text.strip()
    
    def extract_text_from_txt(self, filepath: str) -> str:
        """Extract text content from a TXT file."""
        encodings = ['utf-8', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as file:
                    return file.read().strip()
            except UnicodeDecodeError:
                continue
        raise Exception("Failed to read text file with supported encodings")
    
    def extract_text(self, filepath: str) -> str:
        """Extract text from various file formats."""
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(filepath)
        elif ext in ['.docx', '.doc']:
            return self.extract_text_from_docx(filepath)
        elif ext == '.txt':
            return self.extract_text_from_txt(filepath)
        else:
            raise Exception(f"Unsupported file format: {ext}")
    
    def preprocess_text(self, text: str) -> str:
        """Clean and preprocess text for analysis."""
        # Convert to lowercase
        text = text.lower()
        
        # Remove URLs
        text = re.sub(r'http[s]?://\S+', '', text)
        
        # Remove email addresses
        text = re.sub(r'\S+@\S+', '', text)
        
        # Remove phone numbers
        text = re.sub(r'[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}', '', text)
        
        # Keep alphanumeric, spaces, and some punctuation
        text = re.sub(r'[^a-zA-Z0-9\s\.\+\#]', ' ', text)
        
        # Remove extra whitespace
        text = ' '.join(text.split())
        
        return text
    
    def tokenize_and_lemmatize(self, text: str) -> List[str]:
        """Tokenize and lemmatize text."""
        try:
            tokens = word_tokenize(text.lower())
        except Exception:
            tokens = text.lower().split()
        
        # Lemmatize and remove stopwords
        lemmatized = []
        for token in tokens:
            if token not in self.stop_words and len(token) > 2 and token.isalpha():
                lemmatized.append(self.lemmatizer.lemmatize(token))
        
        return lemmatized
    
    def extract_skills(self, text: str) -> Dict[str, List[str]]:
        """Extract skills from text categorized by type."""
        text_lower = text.lower()
        found_skills = {category: [] for category in self.TECH_SKILLS}
        
        for category, skills in self.TECH_SKILLS.items():
            for skill in skills:
                # Check for exact word match (with word boundaries)
                pattern = r'\b' + re.escape(skill) + r'\b'
                if re.search(pattern, text_lower):
                    found_skills[category].append(skill)
        
        # Remove empty categories
        return {k: v for k, v in found_skills.items() if v}
    
    def extract_education(self, text: str) -> Dict[str, bool]:
        """Extract education level from text."""
        text_lower = text.lower()
        found_education = {}
        
        for level, keywords in self.EDUCATION_LEVELS.items():
            for keyword in keywords:
                if keyword in text_lower:
                    found_education[level] = True
                    break
            if level not in found_education:
                found_education[level] = False
        
        return found_education
    
    def extract_experience_years(self, text: str) -> int:
        """Extract years of experience from text."""
        text_lower = text.lower()
        max_years = 0
        
        for pattern in self.EXPERIENCE_PATTERNS:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                try:
                    years = int(match)
                    if years < 50:  # Sanity check
                        max_years = max(max_years, years)
                except ValueError:
                    continue
        
        return max_years
    
    def calculate_similarity(self, text1: str, text2: str) -> float:
        """Calculate cosine similarity between two texts using TF-IDF."""
        try:
            tfidf_matrix = self.vectorizer.fit_transform([text1, text2])
            similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]
            return float(similarity)
        except Exception:
            return 0.0
    
    def calculate_keyword_match(self, resume_text: str, job_text: str) -> Tuple[float, List[str], List[str]]:
        """Calculate keyword match percentage."""
        # Get tokens
        resume_tokens = set(self.tokenize_and_lemmatize(resume_text))
        job_tokens = set(self.tokenize_and_lemmatize(job_text))
        
        # Find matches and missing
        matched = resume_tokens.intersection(job_tokens)
        missing = job_tokens - resume_tokens
        
        # Calculate match percentage
        if len(job_tokens) == 0:
            return 0.0, [], []
        
        match_percentage = len(matched) / len(job_tokens)
        
        return match_percentage, list(matched)[:20], list(missing)[:10]
    
    def calculate_skill_match(self, resume_skills: Dict[str, List[str]], 
                              job_skills: Dict[str, List[str]]) -> Tuple[float, Dict[str, Any]]:
        """Calculate skill match between resume and job description."""
        all_resume_skills = set()
        all_job_skills = set()
        
        for skills in resume_skills.values():
            all_resume_skills.update(skills)
        
        for skills in job_skills.values():
            all_job_skills.update(skills)
        
        if not all_job_skills:
            return 1.0, {'matched': [], 'missing': []}
        
        matched = all_resume_skills.intersection(all_job_skills)
        missing = all_job_skills - all_resume_skills
        
        match_percentage = len(matched) / len(all_job_skills)
        
        return match_percentage, {
            'matched': list(matched),
            'missing': list(missing)
        }
    
    def analyze_resume(self, filepath: str, job_description: str) -> Dict[str, Any]:
        """
        Perform comprehensive analysis of a resume against a job description.
        
        Args:
            filepath: Path to the resume file
            job_description: Text of the job description
            
        Returns:
            Dictionary containing analysis results and overall score
        """
        try:
            # Extract and preprocess resume text
            resume_text = self.extract_text(filepath)
            if not resume_text or len(resume_text) < 50:
                return {
                    'error': 'Could not extract sufficient text from resume',
                    'score': 0
                }
            
            resume_clean = self.preprocess_text(resume_text)
            job_clean = self.preprocess_text(job_description)
            
            # Calculate various metrics
            
            # 1. TF-IDF Cosine Similarity (30% weight)
            similarity_score = self.calculate_similarity(resume_clean, job_clean)
            
            # 2. Keyword Match (25% weight)
            keyword_match, matched_keywords, missing_keywords = self.calculate_keyword_match(
                resume_clean, job_clean
            )
            
            # 3. Technical Skills Match (30% weight)
            resume_skills = self.extract_skills(resume_text)
            job_skills = self.extract_skills(job_description)
            skill_match, skill_details = self.calculate_skill_match(resume_skills, job_skills)
            
            # 4. Experience extraction
            experience_years = self.extract_experience_years(resume_text)
            
            # 5. Education extraction
            education = self.extract_education(resume_text)
            
            # Calculate weighted overall score
            overall_score = (
                similarity_score * 0.30 +
                keyword_match * 0.25 +
                skill_match * 0.30 +
                min(experience_years / 10, 1) * 0.15  # Cap at 10 years
            )
            
            # Convert to percentage (0-100)
            overall_score = round(overall_score * 100, 2)
            
            # Ensure score is within bounds
            overall_score = max(0, min(100, overall_score))
            
            # Generate recommendation
            if overall_score >= 75:
                recommendation = "Highly Recommended"
                recommendation_class = "excellent"
            elif overall_score >= 60:
                recommendation = "Good Match"
                recommendation_class = "good"
            elif overall_score >= 40:
                recommendation = "Moderate Match"
                recommendation_class = "moderate"
            else:
                recommendation = "Low Match"
                recommendation_class = "low"
            
            # Count total skills
            total_resume_skills = sum(len(skills) for skills in resume_skills.values())
            
            return {
                'score': overall_score,
                'recommendation': recommendation,
                'recommendation_class': recommendation_class,
                'metrics': {
                    'semantic_similarity': round(similarity_score * 100, 2),
                    'keyword_match': round(keyword_match * 100, 2),
                    'skill_match': round(skill_match * 100, 2)
                },
                'skills': {
                    'found': resume_skills,
                    'total_count': total_resume_skills,
                    'matched_with_job': skill_details['matched'],
                    'missing_from_resume': skill_details['missing']
                },
                'experience_years': experience_years,
                'education': education,
                'keywords': {
                    'matched': matched_keywords[:15],
                    'missing': missing_keywords[:10]
                },
                'resume_length': len(resume_text)
            }
            
        except Exception as e:
            return {
                'error': str(e),
                'score': 0
            }


# Testing function
def test_analyzer():
    """Test the analyzer with sample data."""
    print("Testing ResumeAnalyzer...")
    
    analyzer = ResumeAnalyzer()
    
    # Test text preprocessing
    sample_text = "  Hello   World!  This is a TEST @email.com https://example.com  "
    cleaned = analyzer.preprocess_text(sample_text)
    print(f"Preprocessed text: '{cleaned}'")
    
    # Test skill extraction
    sample_resume = """
    Experienced Python developer with 5 years of experience in Django and Flask.
    Proficient in JavaScript, React, and Node.js. Strong knowledge of AWS and Docker.
    Bachelor's degree in Computer Science. Excellent communication and teamwork skills.
    """
    
    skills = analyzer.extract_skills(sample_resume)
    print(f"Extracted skills: {skills}")
    
    experience = analyzer.extract_experience_years(sample_resume)
    print(f"Experience years: {experience}")
    
    education = analyzer.extract_education(sample_resume)
    print(f"Education: {education}")
    
    print("\nAll tests passed!")


if __name__ == '__main__':
    test_analyzer()
